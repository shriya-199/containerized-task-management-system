from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "TaskFlow-AI-INT332-Report-Shriya-Verma.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(85, 85, 85)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(row.cells[idx])


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TaskFlow AI Project Report")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_title_page(doc):
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("TaskFlow AI")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = DARK_BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Advanced Todo Management Application\nUsing React.js and Docker-Based DevOps Workflow")
    r.bold = True
    r.font.size = Pt(16)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Report").bold = True
    p.runs[0].font.size = Pt(16)

    for text in [
        "Submitted in partial fulfilment of the requirements for the award of degree of",
        "BTech - Computer Science and Engineering",
        "LOVELY PROFESSIONAL UNIVERSITY",
        "PHAGWARA, PUNJAB",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(12)
        if "BTech" in text or "LOVELY" in text:
            r.bold = True

    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [2.1, 4.2])
    data = [
        ("Submitted By", "Name of student: Shriya Verma\nRegistration Number: 12323829\nSignature of the student: Shriya Verma"),
        ("Submitted To", "Faculty: Prof. Divya Thakur"),
        ("Course", "INT332 - DevOps"),
        ("Project", "TaskFlow AI - Containerized Task Management System"),
        ("Date", "30 May 2026"),
    ]
    for row, (left, right) in zip(table.rows, data):
        set_cell_shading(row.cells[0], "E8EEF5")
        row.cells[0].paragraphs[0].add_run(left).bold = True
        row.cells[1].paragraphs[0].add_run(right)
    doc.add_page_break()


def add_section(doc, heading, paragraphs):
    doc.add_heading(heading, level=1)
    for text in paragraphs:
        doc.add_paragraph(text)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [2.0, 4.3])
    hdr = table.rows[0].cells
    hdr[0].text = "Component"
    hdr[1].text = "Description"
    for cell in hdr:
        set_cell_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            set_cell_margins(cell)
    doc.add_paragraph()


def add_toc(doc):
    doc.add_heading("3. Table of Contents", level=1)
    entries = [
        "1. Acknowledgement",
        "2. Abstract",
        "3. Table of Contents",
        "4. List of Abbreviations",
        "5. Introduction",
        "6. Literature Review and Technology Background",
        "7. Objectives of the Project",
        "8. Methodology",
        "9. Requirements and Technology Stack",
        "10. System Design and Architecture",
        "11. Implementation Details",
        "12. Docker and DevOps Workflow",
        "13. Testing and Verification",
        "14. Applications and Future Scope",
        "15. Conclusion",
        "16. References",
    ]
    for entry in entries:
        doc.add_paragraph(entry, style="List Number")
    doc.add_page_break()


def build():
    doc = Document()
    style_doc(doc)
    footer(doc.sections[0])
    add_title_page(doc)

    doc.add_heading("Student Declaration", level=1)
    doc.add_paragraph(
        "I, Shriya Verma, 12323829, hereby declare that the work done by me on TaskFlow AI: Advanced Todo Management "
        "Application Using React.js and Docker-Based DevOps Workflow is a record of original work "
        "carried out for the partial fulfilment of the requirements for the award of the degree, "
        "BTech - Computer Science and Engineering."
    )
    doc.add_paragraph("Name of the Student (Registration Number): Shriya Verma (12323829)")
    doc.add_paragraph("Signature of the student: Shriya Verma")
    doc.add_paragraph("Dated: 30 May 2026")
    doc.add_page_break()

    add_section(
        doc,
        "1. Acknowledgement",
        [
            "I would like to express my sincere gratitude to Lovely Professional University for providing me with the opportunity to undertake this project as part of my academic curriculum. The project helped me understand the practical implementation of frontend engineering, containerization, and DevOps workflow automation.",
            "I extend my thanks to my faculty guide and the Computer Science and Engineering department for their continuous guidance, technical feedback, and support throughout the development of this project.",
            "I am also thankful to the open-source communities behind React.js, Vite, Tailwind CSS, Docker, Node.js, GitHub Actions, and related tools. Their documentation and ecosystem made it possible to build a structured, modern, and containerized application.",
            "Lastly, I thank my classmates, peers, and family members for their encouragement and support during the completion of this project.",
            "Shriya Verma\nRegistration Number: 12323829",
        ],
    )
    doc.add_page_break()

    add_section(
        doc,
        "2. Abstract",
        [
            "Modern software systems require not only functional application development but also reliable deployment workflows. DevOps practices, containerization, and automated verification help reduce deployment errors and environment inconsistencies. This project demonstrates these concepts through TaskFlow AI, an advanced todo management application built using React.js and a Docker-based workflow.",
            "TaskFlow AI is a premium task management application inspired by modern productivity platforms. It provides a landing page, authentication screens, dashboard, task management, task details, calendar view, analytics, settings, dark/light theme support, smooth animations, and localStorage persistence. The application uses React Router DOM for navigation, Context API for state management, Tailwind CSS for styling, Framer Motion for animations, Lucide React for icons, and Recharts for analytics visualization.",
            "The project is containerized using Docker. The frontend is built with Node.js inside a multi-stage Dockerfile and served using Nginx. A lightweight Node.js Express backend service is also included to demonstrate service health checks and backend containerization. Docker Compose orchestrates the frontend and backend services, while GitHub Actions validates the Docker Compose configuration, builds images, starts containers, and verifies the application automatically.",
            "Through this project, practical understanding of React frontend architecture, component-based design, local persistence, Docker image creation, Docker Compose service orchestration, and CI/CD automation has been achieved.",
        ],
    )
    doc.add_page_break()
    add_toc(doc)

    doc.add_heading("4. List of Abbreviations", level=1)
    add_key_value_table(
        doc,
        [
            ("API", "Application Programming Interface"),
            ("CI/CD", "Continuous Integration and Continuous Deployment"),
            ("CSS", "Cascading Style Sheets"),
            ("DOM", "Document Object Model"),
            ("SPA", "Single Page Application"),
            ("UI/UX", "User Interface and User Experience"),
            ("VCS", "Version Control System"),
            ("YAML", "Yet Another Markup Language"),
        ],
    )
    doc.add_page_break()

    add_section(
        doc,
        "5. Introduction",
        [
            "Task management applications are widely used by students, professionals, and teams to organize daily work. A good task management system should allow users to create tasks, assign priorities, monitor progress, and review productivity through a clean interface.",
            "TaskFlow AI was developed as a modern React.js application that demonstrates both frontend engineering and DevOps concepts. The project focuses on building a usable, responsive, and visually polished task management system while also showing how the application can be packaged and executed through Docker containers.",
        ],
    )
    doc.add_heading("5.1 Overview of DevOps and Containerization", level=2)
    doc.add_paragraph(
        "DevOps combines development and operations practices to improve software delivery. Containerization packages an application and its dependencies into isolated units called containers. This avoids environment mismatch issues and makes applications easier to run on different machines."
    )
    doc.add_heading("5.2 Project Motivation", level=2)
    add_bullets(
        doc,
        [
            "To build a practical React.js application that can be demonstrated in a classroom environment.",
            "To apply frontend architecture concepts such as reusable components, routing, state management, and responsive UI.",
            "To demonstrate Docker-based application packaging and Docker Compose orchestration.",
            "To include GitHub Actions as an automated verification pipeline.",
        ],
    )
    doc.add_page_break()

    add_section(
        doc,
        "6. Literature Review and Technology Background",
        [
            "React.js is a popular JavaScript library for building component-based user interfaces. Vite improves the frontend development experience by providing fast builds and a simple configuration model. Tailwind CSS supports utility-first styling and helps produce consistent responsive interfaces.",
            "Docker is widely used in software engineering for containerization. It allows developers to build reproducible images and run services in isolated containers. Docker Compose is used when multiple services need to be started together as one application stack.",
        ],
    )
    doc.add_heading("6.1 React and SPA Development", level=2)
    doc.add_paragraph(
        "A Single Page Application loads once in the browser and updates views dynamically without full page reloads. React Router DOM enables multiple routes such as dashboard, tasks, calendar, analytics, and settings within the same frontend application."
    )
    doc.add_heading("6.2 Docker-Based Deployment", level=2)
    doc.add_paragraph(
        "The frontend Dockerfile uses a multi-stage build. The first stage uses Node.js to install dependencies and generate production files. The second stage uses Nginx to serve the generated static files. This approach keeps the final runtime image smaller and more production-oriented."
    )
    doc.add_page_break()

    doc.add_heading("7. Objectives of the Project", level=1)
    add_numbered(
        doc,
        [
            "To design and develop a complete task management application named TaskFlow AI.",
            "To implement modern frontend architecture using React.js, React Router DOM, Context API, and reusable components.",
            "To provide localStorage persistence for tasks, user data, and theme preferences.",
            "To implement dark/light theme, animations, charts, task filtering, search, sorting, and drag-and-drop reordering.",
            "To containerize the frontend and backend services using Docker.",
            "To orchestrate services using Docker Compose.",
            "To automate project verification using GitHub Actions.",
        ],
    )
    doc.add_page_break()

    doc.add_heading("8. Methodology", level=1)
    doc.add_heading("8.1 Requirement Analysis", level=2)
    doc.add_paragraph(
        "The project requirements were divided into functional and DevOps requirements. Functional requirements included task creation, editing, deletion, completion, importance marking, dashboard statistics, calendar view, analytics, and settings. DevOps requirements included Dockerfiles, Docker Compose, container verification, and CI workflow."
    )
    doc.add_heading("8.2 Development Workflow", level=2)
    add_numbered(
        doc,
        [
            "Create the React frontend structure using Vite.",
            "Build reusable components such as sidebar, task card, modal, task form, toast center, and statistic cards.",
            "Implement Context API providers for task state, theme state, user state, and toast notifications.",
            "Add routing for all pages using React Router DOM.",
            "Style the application using Tailwind CSS and responsive design principles.",
            "Add Dockerfile, Nginx configuration, and Docker Compose service definitions.",
            "Verify the build locally and through Docker containers.",
        ],
    )
    doc.add_page_break()

    doc.add_heading("9. Requirements and Technology Stack", level=1)
    doc.add_heading("9.1 Hardware Requirements", level=2)
    add_bullets(
        doc,
        [
            "Processor: Intel i3 or above",
            "RAM: 4 GB minimum, 8 GB recommended",
            "Storage: 2 GB free space for source code, dependencies, Docker images, and containers",
            "Internet connection for dependency and Docker image downloads",
        ],
    )
    doc.add_heading("9.2 Software Requirements", level=2)
    add_key_value_table(
        doc,
        [
            ("Operating System", "Windows 10/11 or Linux-based system"),
            ("Frontend Runtime", "Node.js 18.20.4 for Docker build stage"),
            ("Frontend Library", "React.js with Vite"),
            ("Styling", "Tailwind CSS"),
            ("Containerization", "Docker and Docker Compose"),
            ("CI/CD", "GitHub Actions"),
            ("Editor", "Visual Studio Code"),
        ],
    )
    doc.add_page_break()

    doc.add_heading("10. System Design and Architecture", level=1)
    doc.add_paragraph(
        "The system follows a modular frontend architecture. Pages are separated from components, contexts, hooks, utility functions, and routes. This improves maintainability and makes the project easier to explain and extend."
    )
    add_key_value_table(
        doc,
        [
            ("Frontend", "React.js SPA served by Nginx in Docker"),
            ("State Layer", "Context API for tasks, theme, user, and toast notifications"),
            ("Persistence", "Browser localStorage for tasks, theme, and user profile"),
            ("Backend", "Node.js Express service used for health checks and backend container demonstration"),
            ("Orchestration", "Docker Compose starts frontend and backend services together"),
        ],
    )
    doc.add_heading("10.1 Application Modules", level=2)
    add_bullets(
        doc,
        [
            "Landing module introduces the application and provides navigation to login/signup/demo.",
            "Dashboard module displays summary statistics, recent tasks, upcoming tasks, and charts.",
            "Task module supports add, edit, delete, complete, restore, important, search, filter, and sort.",
            "Calendar module shows tasks by due date in a monthly layout.",
            "Analytics module visualizes completion rate, weekly productivity, categories, and priorities.",
            "Settings module manages profile, theme, notifications, and danger-zone actions.",
        ],
    )
    doc.add_page_break()

    doc.add_heading("11. Implementation Details", level=1)
    doc.add_heading("11.1 Folder Structure", level=2)
    add_key_value_table(
        doc,
        [
            ("src/components", "Reusable UI components such as TaskCard, TaskForm, Modal, ToastCenter, StatCard, and charts"),
            ("src/pages", "Application pages such as Dashboard, Task List, Calendar, Analytics, Login, Signup, and Settings"),
            ("src/context", "Context providers for tasks, theme, user, and toast notifications"),
            ("src/hooks", "Custom hooks such as useLocalStorage, useTheme, and useTasks"),
            ("src/routes", "Central route configuration using React Router DOM"),
            ("src/utils", "Helper functions for dates, task creation, sorting, and productivity calculations"),
        ],
    )
    doc.add_heading("11.2 Core Features", level=2)
    add_bullets(
        doc,
        [
            "Task CRUD: Users can add, edit, delete, and view task details.",
            "Status management: Tasks can be marked completed or restored to pending.",
            "Important tasks: Starred tasks appear in the Important section.",
            "Search and filters: Users can search tasks and filter by completed, pending, or important status.",
            "Sorting: Tasks can be sorted by date, priority, or alphabetical order.",
            "Theme: Dark and light modes are stored in localStorage.",
            "Analytics: Recharts is used to display productivity and task distribution charts.",
            "Animations: Framer Motion provides page transitions, card hover effects, and modal animations.",
        ],
    )
    doc.add_page_break()

    doc.add_heading("12. Docker and DevOps Workflow", level=1)
    doc.add_heading("12.1 Dockerfile Implementation", level=2)
    doc.add_paragraph(
        "The frontend Dockerfile uses two stages. In the build stage, Node.js installs dependencies and runs the Vite production build. In the runtime stage, Nginx serves the generated files from the dist directory. A custom Nginx configuration supports React Router by redirecting unknown routes to index.html."
    )
    add_key_value_table(
        doc,
        [
            ("Frontend Container", "Serves the React production build using Nginx on container port 80"),
            ("Backend Container", "Runs the Node.js Express backend on container port 3000"),
            ("Docker Compose", "Maps frontend to localhost:8081 and backend to localhost:3000"),
            ("Health Check", "Backend exposes GET /health for verification"),
        ],
    )
    doc.add_heading("12.2 Docker Compose Services", level=2)
    add_bullets(
        doc,
        [
            "daily-task-manager-frontend: React/Nginx frontend container.",
            "daily-task-manager-backend: Node.js Express backend container.",
            "containerized-task-management-system-frontend:latest: frontend Docker image.",
            "containerized-task-management-system-backend:latest: backend Docker image.",
        ],
    )
    doc.add_heading("12.3 GitHub Actions Workflow", level=2)
    doc.add_paragraph(
        "The GitHub Actions workflow runs on pushes and pull requests to the main branch. It validates Docker Compose configuration, builds images, starts containers, verifies backend health, verifies the TaskFlow AI frontend, shows container status, and tears down the stack."
    )
    doc.add_page_break()

    doc.add_heading("13. Testing and Verification", level=1)
    add_key_value_table(
        doc,
        [
            ("Frontend Build", "npm run build successfully generates production files"),
            ("Docker Build", "docker compose up --build -d builds and starts containers"),
            ("Frontend Verification", "http://localhost:8081 returns the TaskFlow AI application"),
            ("Route Verification", "http://localhost:8081/app/dashboard works due to Nginx SPA fallback"),
            ("Backend Verification", "http://localhost:3000/health returns healthy status"),
        ],
    )
    doc.add_heading("13.1 Demo Commands", level=2)
    for command in [
        'cd "C:\\Users\\hp\\OneDrive\\Desktop\\containerized-task-management-system"',
        "docker-compose up --build -d",
        "docker-compose ps",
        "docker-compose logs -f frontend",
        "docker-compose logs -f backend",
        "docker-compose down",
    ]:
        p = doc.add_paragraph()
        r = p.add_run(command)
        r.font.name = "Consolas"
        r.font.size = Pt(10)
    doc.add_page_break()

    doc.add_heading("14. Applications and Future Scope", level=1)
    doc.add_heading("14.1 Current Applications", level=2)
    add_bullets(
        doc,
        [
            "Personal daily task management for students and individuals.",
            "Academic demonstration of React frontend architecture.",
            "DevOps demonstration using Docker, Docker Compose, and GitHub Actions.",
            "Portfolio project for frontend and DevOps learning.",
        ],
    )
    doc.add_heading("14.2 Future Enhancements", level=2)
    add_bullets(
        doc,
        [
            "Add real user authentication with JWT and backend database storage.",
            "Replace localStorage with MySQL or MongoDB for multi-user persistence.",
            "Add team collaboration, shared workspaces, and task assignment.",
            "Add email reminders and push notifications.",
            "Add automated unit tests and end-to-end tests.",
            "Deploy the application to a cloud platform using CI/CD.",
        ],
    )
    doc.add_page_break()

    add_section(
        doc,
        "15. Conclusion",
        [
            "The TaskFlow AI project successfully demonstrates the development of a modern todo management application using React.js and supporting frontend technologies. The application provides a polished user interface, routing, task management, localStorage persistence, theme management, animations, and analytics.",
            "From the DevOps perspective, the project demonstrates Docker image building, Nginx-based frontend serving, backend containerization, Docker Compose orchestration, and GitHub Actions automation. These practices reflect important industry workflows used to build and verify applications in a reliable manner.",
            "The project helped strengthen understanding of frontend architecture, reusable components, state management, responsive design, containerization, and CI/CD. It can be further extended into a full production-grade multi-user task management system by adding authentication, database persistence, cloud deployment, and automated testing.",
        ],
    )

    doc.add_heading("16. References", level=1)
    add_numbered(
        doc,
        [
            "React Documentation: https://react.dev/",
            "Vite Documentation: https://vitejs.dev/",
            "Tailwind CSS Documentation: https://tailwindcss.com/",
            "Docker Documentation: https://docs.docker.com/",
            "GitHub Actions Documentation: https://docs.github.com/actions",
            "React Router Documentation: https://reactrouter.com/",
            "Framer Motion Documentation: https://www.framer.com/motion/",
            "Recharts Documentation: https://recharts.org/",
        ],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
